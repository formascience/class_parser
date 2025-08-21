from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


class DocxWriter:
    """Generate a professionally formatted .docx from a Course-like object.

    The provided ``course`` object is expected to expose at least the following
    attributes used by this writer: ``name``, ``course_title``, ``level``,
    ``block``, ``semester``, ``subject``, and ``content`` (with hierarchical
    sections structure matching the models in ``src/models.py``).
    """

    def __init__(
        self,
        template_path: str = "volume/fs_template.docx",
        title_font: str = "Manrope",
        title_size: int = 14,
        title_color: tuple[int, int, int] = (0, 0, 0),
        content_font: str = "Arial",
        content_size: int = 10,
        table_font: str = "Arial",
        table_size: int = 11,
        table_color: tuple[int, int, int] = (0, 0, 0),
    ) -> None:
        self.template_path = template_path

        # Title formatting
        self.title_font = title_font
        self.title_size = title_size
        self.title_color = title_color

        # Content formatting
        self.content_font = content_font
        self.content_size = content_size

        # Table formatting
        self.table_font = table_font
        self.table_size = table_size
        self.table_color = table_color

    def fill_template(self, course: Any, output_path: Optional[str] = None) -> str:
        """Fill the Word template with course metadata and content.

        Returns the path to the saved .docx file.
        """
        doc = Document(self.template_path)

        # Page-level spacing: ensure a consistent gap after page header on all pages
        #for section in doc.sections:
            # section.header_distance = Pt(24)  # distance from page top to header
            # section.top_margin = Pt(108)  # 1.5 inches to keep body below header
            # section.left_margin = Pt(72)
            # section.right_margin = Pt(72)
            # section.bottom_margin = Pt(72)

        self._fill_title(doc, course)
        self._fill_table(doc, course)
        self._fill_header(doc, course)

        if getattr(course, "content", None):
            self._write_content(doc, course.content)

        if not output_path:
            saved_name = (course.name or "course").lower().replace(" ", "_")
            output_path = f"volume/artifacts/{saved_name}_filled.docx"

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    def _fill_title(self, doc: Any, course: Any) -> None:
        for paragraph in doc.paragraphs:
            if "TITRE DU COURS" in paragraph.text:
                course_title = f"Chapitre {course.chapter} : {course.course_title}".upper()
                paragraph.text = course_title
                for run in paragraph.runs:
                    run.font.name = self.title_font
                    run.font.size = Pt(14)  # Course title font size 14
                    run.font.bold = False
                    r, g, b = self.title_color
                    run.font.color.rgb = RGBColor(r, g, b)
                break

    def _fill_table(self, doc: Any, course: Any) -> None:
        if len(doc.tables) == 0:
            return

        table = doc.tables[0]
        cells_data = [
            (0, 0, f"BLOC {getattr(course, 'block', None) or 'SANTE'}"),
            (1, 0, getattr(course, "subject", "")),
            (0, 1, "Auteur : RonéoAI"),
            (1, 1, f"Date : {datetime.now().strftime('%Y-%m-%d')}")
        ]

        for row, col, text in cells_data:
            cell = table.rows[row].cells[col]
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = self.table_font
                    run.font.size = Pt(self.table_size)
                    r, g, b = self.table_color
                    run.font.color.rgb = RGBColor(r, g, b)

    def _fill_header(self, doc: Any, course: Any) -> None:
        for paragraph in doc.paragraphs:
            if "L1.SpS" in paragraph.text:
                new_level = f"{getattr(course, 'level', None) or 'L1'}.{getattr(course, 'semester', None) or 'S1'}"
                paragraph.text = paragraph.text.replace("L1.SpS", new_level)
                break

    def _write_content(self, doc: Any, content: Any) -> None:
        for i, section in enumerate(content.sections, start=1):
            clean_title = self._strip_existing_numbering(section.title)
            roman_title = f"{self._to_roman(i)}. {clean_title}".upper()
            heading1 = doc.add_paragraph(roman_title, style="Heading 1")

            heading1_format = heading1.paragraph_format
            heading1_format.space_before = Pt(12)
            heading1_format.space_after = Pt(12)

            for run in heading1.runs:
                run.font.name = self.title_font  # Level 0 uses title font
                run.font.size = Pt(12)  # Level 0 sections font size 14
                run.font.bold = False
                
                # Leaf green
                run.font.color.rgb = RGBColor(143, 150, 78)

            if section.content:
                self._write_content_items(doc, section.content)

            self._write_subsections(doc, section.subsections, i)

    def _write_subsections(
        self, doc: Any, subsections: list[Any], parent_num: int | str, level: int = 1
    ) -> None:
        for j, subsection in enumerate(subsections, start=1):
            clean_title = self._strip_existing_numbering(subsection.title)
            if level == 1:
                subtitle = f"{j}. {clean_title}"
                heading = doc.add_paragraph(subtitle, style="Heading 2")
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
                for run in heading.runs:
                    run.font.name = self.table_font
                    run.font.size = Pt(11)  # Other section titles font size 12
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif level == 2:
                subtitle = f"{parent_num}.{j} {clean_title}"
                heading = doc.add_paragraph(subtitle, style="Heading 3")
                heading.paragraph_format.space_before = Pt(10)
                heading.paragraph_format.space_after = Pt(6)
                for run in heading.runs:
                    run.font.name = self.table_font
                    run.font.size = Pt(11)  # Other section titles font size 12
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                subtitle = f"{parent_num}.{j} {clean_title}"
                heading = doc.add_paragraph(subtitle)
                heading.paragraph_format.space_before = Pt(8)
                heading.paragraph_format.space_after = Pt(4)
                for run in heading.runs:
                    run.font.name = self.table_font
                    run.font.size = Pt(11)  # Other section titles font size 12
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)

            if getattr(subsection, "content", None):
                self._write_content_items(doc, subsection.content)

            if getattr(subsection, "subsections", None):
                self._write_subsections(
                    doc, subsection.subsections, f"{parent_num}.{j}", level + 1
                )

    def _write_content_items(self, doc: Any, content_items: list[str]) -> None:
        """Process content items with structured formatting support"""
        for content_item in content_items:
            if not content_item or not content_item.strip():
                continue
                
            content_item = content_item.strip()

            # If a bullet marker appears anywhere, treat the entire string as a list
            if '•' in content_item:
                bullet_items = self._split_bullet_items(content_item)
                if bullet_items:
                    for bi in bullet_items:
                        p = doc.add_paragraph()
                        self._add_bullet_content(p, bi)
                        self._format_list_paragraph(p)
                    continue  # handled

            # If a numbered-marker pattern appears anywhere, split into items
            if re.search(r'(?<!\S)\d+\.\s', content_item):
                numbered_items = self._split_numbered_items(content_item)
                if numbered_items:
                    for number, text in numbered_items:
                        p = doc.add_paragraph()
                        self._add_numbered_content(p, number, text)
                        self._format_list_paragraph(p)
                    continue  # handled

            # Fallback single-line handlers
            if content_item.startswith('• '):
                text = content_item[2:].strip()
                p = doc.add_paragraph()
                self._add_bullet_content(p, text)
                self._format_list_paragraph(p)
            elif re.match(r'^\d+\.\s', content_item):
                match = re.match(r'^(\d+)\.\s(.+)', content_item, re.DOTALL)
                if match:
                    p = doc.add_paragraph()
                    self._add_numbered_content(p, match.group(1), match.group(2).strip())
                    self._format_list_paragraph(p)
                else:
                    p = doc.add_paragraph(content_item)
                    self._format_regular_paragraph(p)
            elif content_item.startswith('- ') or content_item.startswith('* '):
                text = content_item[2:].strip()
                p = doc.add_paragraph()
                self._add_bullet_content(p, text)
                self._format_list_paragraph(p)
            else:
                p = doc.add_paragraph(content_item)
                self._format_regular_paragraph(p)

    def _split_bullet_items(self, text: str) -> list[str]:
        """Split a paragraph containing multiple '•' into separate bullet items.

        Example: "• A • B • C" -> ["A", "B", "C"]
        """
        if not text:
            return []
        # Normalize spaces around bullets and split
        parts = re.split(r"\s*•\s+", text.strip())
        # If the text started with a bullet, the first item will be empty; drop empties
        items = [p.strip() for p in parts if p.strip()]
        return items

    def _split_numbered_items(self, text: str) -> list[tuple[str, str]]:
        """Split a paragraph containing multiple 'n. ' patterns into numbered items.

        Returns list of tuples (number, text).
        """
        items: list[tuple[str, str]] = []
        if not text:
            return items

        pattern = re.compile(r"(?<!\S)(\d+)\.\s")  # number dot space at start or after whitespace
        pos = 0
        last_num: Optional[str] = None
        for m in pattern.finditer(text):
            if last_num is not None:
                items.append((last_num, text[pos:m.start()].strip()))
            last_num = m.group(1)
            pos = m.end()
        if last_num is not None:
            items.append((last_num, text[pos:].strip()))
        return [(n, t) for (n, t) in items if t]

    def _add_bullet_content(self, paragraph, text: str) -> None:
        """Add bullet point content to a paragraph with manual formatting"""
        # Add bullet symbol as first run
        bullet_run = paragraph.add_run("• ")
        bullet_run.font.name = self.content_font
        bullet_run.font.size = Pt(self.content_size)
        bullet_run.font.bold = False
        
        # Add the text content as second run
        text_run = paragraph.add_run(text)
        text_run.font.name = self.content_font
        text_run.font.size = Pt(self.content_size)
        text_run.font.bold = False
        
        # No indentation - bullet points aligned with regular text
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)

    def _add_numbered_content(self, paragraph, number: str, text: str) -> None:
        """Add numbered list content to a paragraph with manual formatting"""
        # Add number as first run
        number_run = paragraph.add_run(f"{number}. ")
        number_run.font.name = self.content_font
        number_run.font.size = Pt(self.content_size)
        number_run.font.bold = False
        
        # Add the text content as second run
        text_run = paragraph.add_run(text)
        text_run.font.name = self.content_font
        text_run.font.size = Pt(self.content_size)
        text_run.font.bold = False
        
        # No indentation - numbered lists aligned with regular text
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)

    def _format_list_paragraph(self, paragraph) -> None:
        """Apply consistent formatting to list paragraphs"""
        # Lists typically have tighter spacing
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _format_regular_paragraph(self, paragraph) -> None:
        """Apply formatting to regular paragraphs"""
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = self.content_font
            run.font.size = Pt(self.content_size)

    def _strip_existing_numbering(self, title: str) -> str:
        """Strip common numbering patterns from section titles.
        
        Handles patterns like:
        - "1. Title" -> "Title"
        - "1.1 Title" -> "Title" 
        - "1.1.1 Title" -> "Title"
        - "1.2.4.4.4. Title" -> "Title"
        - "I. Title" -> "Title"
        - "A. Title" -> "Title"
        - "a) Title" -> "Title"
        - "(1) Title" -> "Title"
        """
        # Remove leading/trailing whitespace
        title = title.strip()
        
        # Pattern for various numbering formats at the start of the title
        patterns = [
            r'^\d+(?:\.\d+)*\.?\s*',  # "1.", "1.1", "1.1.1", "1.2.4.4.4.", etc. (any length decimal numbering)
            r'^[IVX]+\.\s*',          # "I. ", "IV. ", "XII. " (Roman numerals)
            r'^[A-Z]\.\s*',           # "A. ", "B. "
            r'^[a-z]\)\s*',           # "a) ", "b) "
            r'^\([0-9]+\)\s*',        # "(1) ", "(123) "
            r'^\([a-z]\)\s*',         # "(a) ", "(b) "
            r'^\([A-Z]\)\s*',         # "(A) ", "(B) "
            r'^-\s*',                 # "- " (bullet point)
            r'^\*\s*',                # "* " (asterisk bullet)
        ]
        
        # Apply each pattern and return the first match that strips something
        for pattern in patterns:
            stripped = re.sub(pattern, '', title)
            if stripped != title:
                return stripped.strip()
        
        # If no pattern matched, return original title
        return title

    def _to_roman(self, num: int) -> str:
        vals = [10, 9, 5, 4, 1]
        syms = ["X", "IX", "V", "IV", "I"]
        roman = ""
        for val, sym in zip(vals, syms):
            count = num // val
            roman += sym * count
            num -= val * count
        return roman


