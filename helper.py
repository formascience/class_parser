def print_readable_format(data, title="DOCUMENT STRUCTURE", max_paragraph_length=150):
    """
    Universal function to print Pydantic models in a readable format.
    Handles CourseDraft, ContentSection, OutlineWithSlides, and CourseDocument.
    
    Args:
        data: The Pydantic model instance to print
        title: Header title for the printout
        max_paragraph_length: Max length before truncating paragraphs
    """
    print("=" * 80)
    print(f"{title.center(80)}")
    print("=" * 80)
    print()
    
    # Handle different data types
    if hasattr(data, 'metadata'):
        # CourseDocument
        _print_course_document_readable(data, max_paragraph_length)
    elif hasattr(data, 'outline') and hasattr(data, 'slide_map'):
        # CourseDraft or OutlineWithSlides
        if hasattr(data.outline[0] if data.outline else None, 'paragraphs'):
            _print_course_draft_readable(data, max_paragraph_length)
        else:
            _print_outline_with_slides_readable(data)
    elif isinstance(data, list):
        # List of ContentSections
        _print_content_sections_readable(data, max_paragraph_length)
    else:
        # Single ContentSection or unknown
        _print_content_section_readable(data, max_paragraph_length)


def _print_course_document_readable(course_doc, max_paragraph_length):
    """Print CourseDocument with metadata and sections."""
    metadata = course_doc.metadata
    print("📚 METADATA:")
    print(f"   📖 Course: {metadata.course_title}")
    print(f"   👨‍🏫 Author: {metadata.author}")
    print(f"   🏢 Bloc: {metadata.bloc}")
    print(f"   📋 UE Code: {metadata.ue_code}")
    print(f"   📅 Date: {metadata.date}")
    print()
    
    print("📖 CONTENT STRUCTURE:")
    print()
    
    for i, section in enumerate(course_doc.sections, 1):
        _print_section_recursive(section, level=1, section_number=i, max_paragraph_length=max_paragraph_length)


def _print_course_draft_readable(course_draft, max_paragraph_length):
    """Print CourseDraft with outline and slide mapping."""
    print("📚 COURSE DRAFT CONTENT:")
    print()
    
    # Print outline with content
    for i, section in enumerate(course_draft.outline, 1):
        _print_content_section_recursive(section, level=1, section_number=i, max_paragraph_length=max_paragraph_length)
    
    # Print slide mapping summary
    print("\n" + "─" * 60)
    print("📄 SLIDE MAPPING SUMMARY")
    print("─" * 60)
    
    # Group slides by path for better readability
    from collections import defaultdict
    path_to_slides = defaultdict(list)
    
    for mapping in course_draft.slide_map:
        path_key = " → ".join(mapping.path) if mapping.path else "(unmapped)"
        path_to_slides[path_key].append(mapping.slide)
    
    for path, slides in sorted(path_to_slides.items()):
        slides_sorted = sorted(slides)
        slides_str = _format_slide_ranges(slides_sorted)
        print(f"  📑 {slides_str:>12} : {path}")


def _print_outline_with_slides_readable(outline_data):
    """Print OutlineWithSlides structure."""
    print("📚 OUTLINE WITH SLIDE REFERENCES:")
    print()
    
    # Build path to slides mapping
    from collections import defaultdict
    path_to_slides = defaultdict(list)
    
    for mapping in outline_data.slide_map:
        path_key = tuple(mapping.path)
        path_to_slides[path_key].append(mapping.slide)
    
    # Print hierarchical outline
    for i, section in enumerate(outline_data.outline, 1):
        _print_outline_section_recursive(section, path=(), path_to_slides=path_to_slides, level=1, section_number=i)


def _print_content_sections_readable(sections, max_paragraph_length):
    """Print a list of ContentSections."""
    print("📚 CONTENT SECTIONS:")
    print()
    
    for i, section in enumerate(sections, 1):
        _print_content_section_recursive(section, level=1, section_number=i, max_paragraph_length=max_paragraph_length)


def _print_content_section_readable(section, max_paragraph_length):
    """Print a single ContentSection."""
    print("📚 CONTENT SECTION:")
    print()
    _print_content_section_recursive(section, level=1, section_number=1, max_paragraph_length=max_paragraph_length)


def _print_section_recursive(section, level=1, section_number=None, max_paragraph_length=150):
    """Recursively print Section (from CourseDocument)."""
    indent = "  " * (level - 1)
    
    # Choose appropriate symbols based on depth
    symbols = ["🔵", "🔸", "▪️", "•", "◦"]
    symbol = symbols[min(level-1, len(symbols)-1)]
    
    title_prefix = f"{section_number}. " if section_number and level == 1 else ""
    print(f"{indent}{symbol} {title_prefix}{section.title}")
    
    # Print paragraphs
    if hasattr(section, 'paragraphs') and section.paragraphs:
        for paragraph in section.paragraphs:
            wrapped = _wrap_text(paragraph, max_paragraph_length, indent)
            print(f"{indent}   📝 {wrapped}")
    
    # Print subsections
    if hasattr(section, 'subsections') and section.subsections:
        for subsection in section.subsections:
            _print_section_recursive(subsection, level + 1, max_paragraph_length=max_paragraph_length)
    
    if level == 1:
        print()


def _print_content_section_recursive(section, level=1, section_number=None, max_paragraph_length=150):
    """Recursively print ContentSection."""
    indent = "  " * (level - 1)
    
    symbols = ["🔵", "🔸", "▪️", "•", "◦"]
    symbol = symbols[min(level-1, len(symbols)-1)]
    
    title_prefix = f"{section_number}. " if section_number and level == 1 else ""
    print(f"{indent}{symbol} {title_prefix}{section.title}")
    
    # Print paragraphs
    if section.paragraphs:
        for paragraph in section.paragraphs:
            wrapped = _wrap_text(paragraph, max_paragraph_length, indent)
            print(f"{indent}   📝 {wrapped}")
    
    # Print subsections
    if section.subsections:
        for subsection in section.subsections:
            _print_content_section_recursive(subsection, level + 1, max_paragraph_length=max_paragraph_length)
    
    if level == 1:
        print()


def _print_outline_section_recursive(section, path=(), path_to_slides=None, level=1, section_number=None):
    """Recursively print Section with slide references."""
    indent = "  " * (level - 1)
    new_path = path + (section.title,)
    
    # Get slides for this path
    slides = path_to_slides.get(new_path, []) if path_to_slides else []
    slide_str = f"  📑 {_format_slide_ranges(sorted(slides))}" if slides else ""
    
    symbols = ["🔵", "🔸", "▪️", "•", "◦"]
    symbol = symbols[min(level-1, len(symbols)-1)]
    
    title_prefix = f"{section_number}. " if section_number and level == 1 else ""
    print(f"{indent}{symbol} {title_prefix}{section.title}{slide_str}")
    
    # Print subsections
    if hasattr(section, 'subsections') and section.subsections:
        for subsection in section.subsections:
            _print_outline_section_recursive(subsection, new_path, path_to_slides, level + 1)
    
    if level == 1:
        print()


def _wrap_text(text, max_length, indent=""):
    """Wrap text to multiple lines when it exceeds max_length."""
    if len(text) <= max_length:
        return text
    
    import textwrap
    # Wrap the text and join with newlines and proper indentation
    wrapped_lines = textwrap.wrap(text, width=max_length, break_long_words=False, break_on_hyphens=False)
    
    # First line gets the emoji, subsequent lines get proper indentation without emoji
    result = wrapped_lines[0]  # First line without extra formatting
    if len(wrapped_lines) > 1:
        # Add subsequent lines with proper indentation (matching the "   📝 " spacing)
        continuation_indent = indent + "      "  # 6 spaces to align with emoji + space
        for line in wrapped_lines[1:]:
            result += f"\n{continuation_indent}{line}"
    
    return result


def _format_slide_ranges(slides):
    """Format a list of slide numbers into ranges (e.g., [1,2,3,5,6] -> "1-3, 5-6")."""
    if not slides:
        return ""
    
    from itertools import groupby
    ranges = []
    for k, g in groupby(enumerate(slides), lambda x: x[1] - x[0]):
        group = list(map(lambda x: x[1], g))
        if len(group) > 1:
            ranges.append(f"{group[0]}-{group[-1]}")
        else:
            ranges.append(str(group[0]))
    
    return ", ".join(ranges)


def export_course_draft_to_docx(course_draft, output_path, title="Course Draft"):
    """
    Export CourseDraft to a DOCX file with hierarchical structure.
    
    Args:
        course_draft: CourseDraft instance with outline and slide_map
        output_path: Path where to save the DOCX file
        title: Document title (default: "Course Draft")
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.shared import OxmlElement, qn
        from collections import defaultdict
    except ImportError:
        raise ImportError("python-docx is required. Install it with: pip install python-docx")
    
    # Create document
    doc = Document()
    
    # Add title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a separator
    doc.add_paragraph()
    
    # Process outline content
    _add_course_content_to_docx(doc, course_draft.outline)
    
    # Add slide mapping section
    _add_slide_mapping_to_docx(doc, course_draft.slide_map)
    
    # Save document
    doc.save(output_path)
    print(f"📄 Course draft exported to: {output_path}")


def _add_course_content_to_docx(doc, outline):
    """Add course content sections to the DOCX document."""
    content_heading = doc.add_heading("📚 Course Content", level=1)
    
    for i, section in enumerate(outline, 1):
        _add_content_section_to_docx(doc, section, level=2, section_number=i)


def _add_content_section_to_docx(doc, section, level=2, section_number=None):
    """Recursively add ContentSection to DOCX document."""
    # Determine heading level (Word supports up to 9 heading levels)
    heading_level = min(level, 9)
    
    # Create section title
    title_prefix = f"{section_number}. " if section_number and level == 2 else ""
    section_heading = doc.add_heading(f"{title_prefix}{section.title}", level=heading_level)
    
    # Add paragraphs
    if section.paragraphs:
        for paragraph_text in section.paragraphs:
            paragraph = doc.add_paragraph()
            # Add a bullet point for each paragraph
            paragraph.style = 'List Bullet'
            paragraph.add_run(paragraph_text)
    
    # Add subsections
    if section.subsections:
        for subsection in section.subsections:
            _add_content_section_to_docx(doc, subsection, level + 1)


def _add_slide_mapping_to_docx(doc, slide_map):
    """Add slide mapping section to the DOCX document."""
    # Add section separator
    doc.add_page_break()
    
    # Add slide mapping heading
    mapping_heading = doc.add_heading("📄 Slide Mapping", level=1)
    
    # Group slides by path for better readability
    from collections import defaultdict
    path_to_slides = defaultdict(list)
    
    for mapping in slide_map:
        path_key = " → ".join(mapping.path) if mapping.path else "(unmapped)"
        path_to_slides[path_key].append(mapping.slide)
    
    # Create a table for better formatting
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Set table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Slides'
    hdr_cells[1].text = 'Section Path'
    
    # Add mapping data
    for path, slides in sorted(path_to_slides.items()):
        row_cells = table.add_row().cells
        slides_sorted = sorted(slides)
        slides_str = _format_slide_ranges(slides_sorted)
        row_cells[0].text = slides_str
        row_cells[1].text = path
    
    # Add some spacing after the table
    doc.add_paragraph()
    
    # Add summary information
    summary_paragraph = doc.add_paragraph()
    summary_paragraph.add_run("Summary: ").bold = True
    summary_paragraph.add_run(f"Total slides mapped: {len(slide_map)}, "
                             f"Unique sections: {len(path_to_slides)}")


# Convenience function for common usage
def export_course_draft(course_draft, filename=None):
    """
    Convenience function to export course draft with default filename.
    
    Args:
        course_draft: CourseDraft instance
        filename: Optional custom filename (default: "course_draft.docx")
    """
    if filename is None:
        filename = "course_draft.docx"
    
    export_course_draft_to_docx(course_draft, filename)
    return filename
