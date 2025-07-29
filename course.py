from typing import List, Optional
from pydantic import BaseModel, Field
import textwrap

class SlideMapping(BaseModel):
    """Maps a slide number to its position in the outline hierarchy"""
    slide_number: int
    section_path: List[str]  # ["Chapter 1", "Introduction", ...]

class ContentSection(BaseModel):
    """A section that can contain both outline structure and content"""
    title: str
    content: List[str] = Field(default_factory=list)  # paragraphs (empty in phase 1, filled in phase 2)
    subsections: List["ContentSection"] = Field(default_factory=list)
    
    def _print_section(self, indent_level: int = 0) -> str:
        """Helper method to print section outline with proper indentation (for phase 1)"""
        indent = "  " * indent_level
        result = f"{indent}{self.title}\n"
        
        for subsection in self.subsections:
            result += subsection._print_section(indent_level + 1)
        
        return result
    
    def _print_content(self, indent_level: int = 0, max_chars_per_line: int = 80) -> str:
        """Helper method to print content section with proper indentation and text wrapping (for phase 2)"""
        indent = "  " * indent_level
        content_indent = "  " * (indent_level + 1)
        result = f"{indent}{self.title}\n"
        
        # Calculate available width for content (subtract indentation)
        available_width = max_chars_per_line - len(content_indent)
        if available_width < 20:  # Minimum reasonable width
            available_width = 20
        
        # Print content paragraphs with text wrapping and additional indentation
        for paragraph in self.content:
            wrapped_lines = textwrap.fill(
                paragraph, 
                width=available_width,
                subsequent_indent=""
            ).split('\n')
            
            for line in wrapped_lines:
                result += f"{content_indent}{line}\n"
        
        # Add spacing between content and subsections if both exist
        if self.content and self.subsections:
            result += "\n"
        
        # Print subsections recursively
        for subsection in self.subsections:
            result += subsection._print_content(indent_level + 1, max_chars_per_line)
        
        return result

ContentSection.model_rebuild()

class Content(BaseModel):
    """Course structure that works for both outline (phase 1) and content (phase 2)"""
    sections: List[ContentSection]
    slide_mappings: List[SlideMapping]
    
    def print_outline(self) -> str:
        """Print the hierarchical outline with indentation (for phase 1)"""
        if not self.sections:
            return "No outline available"
        
        result = "Course Outline:\n"
        result += "=" * 50 + "\n"
        
        for section in self.sections:
            result += section._print_section()
        
        return result.rstrip()  # Remove trailing newline
    
    def print_content(self, max_chars_per_line: int = 80) -> str:
        """Print the hierarchical content with indentation and text wrapping (for phase 2)"""
        if not self.sections:
            return "No content available"
        
        result = "Course Content:\n"
        result += "=" * 50 + "\n"
        
        for section in self.sections:
            result += section._print_content(max_chars_per_line=max_chars_per_line)
        
        return result.rstrip()  # Remove trailing newline

class Course(BaseModel):
    """Main class that represents a complete course with metadata and content"""
    # Course metadata
    name: str  # Course name (e.g., "Introduction to Machine Learning")
    subject: str  # Subject area (e.g., "Computer Science", "Mathematics")
    year: Optional[int] = None  # Academic year (e.g., 2024)
    professor: Optional[str] = None  # Professor name
    
    # Course structure and content (unified for both phases)
    content: Optional[Content] = None  # Works for both outline (phase 1) and content (phase 2)
    
    # Additional metadata
    total_slides: Optional[int] = None
    created_at: Optional[str] = None  # ISO timestamp
    
    def print_outline(self) -> str:
        """Print the course outline with proper indentation (for phase 1)"""
        if not self.content:
            return f"No outline available for course: {self.name}"
        
        header = f"Course: {self.name}\n"
        if self.subject:
            header += f"Subject: {self.subject}\n"
        if self.year:
            header += f"Year: {self.year}\n"
        if self.professor:
            header += f"Professor: {self.professor}\n"
        header += "\n"
        
        return header + self.content.print_outline()
    
    def print_content(self, max_chars_per_line: int = 80) -> str:
        """Print the course content with proper indentation and text wrapping (for phase 2)"""
        if not self.content:
            return f"No content available for course: {self.name}"
        
        header = f"Course: {self.name}\n"
        if self.subject:
            header += f"Subject: {self.subject}\n"
        if self.year:
            header += f"Year: {self.year}\n"
        if self.professor:
            header += f"Professor: {self.professor}\n"
        header += "\n"
        
        return header + self.content.print_content(max_chars_per_line)
    
    def export_to_docx(self, output_path: str = None, include_course_info: bool = True) -> str:
        """
        Export the course content to a DOCX file.
        
        Args:
            output_path: Path where to save the DOCX file (default: "{course_name}.docx")
            include_course_info: Whether to include course metadata in the document
            
        Returns:
            str: Path to the saved DOCX file
        """
        if not self.content:
            raise ValueError(f"No content available for course: {self.name}")
        
        # Import required modules
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from collections import defaultdict
        except ImportError:
            raise ImportError("python-docx is required. Install it with: pip install python-docx")
        
        # Generate default filename if not provided
        if output_path is None:
            safe_name = "".join(c for c in self.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            output_path = f"{safe_name.replace(' ', '_')}.docx"
        
        # Create title with course metadata
        title_parts = [self.name]
        if include_course_info:
            if self.subject:
                title_parts.append(f"({self.subject})")
            if self.year:
                title_parts.append(f"- {self.year}")
            if self.professor:
                title_parts.append(f"- Prof. {self.professor}")
        
        title = " ".join(title_parts)
        
        # Create document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add separator
        doc.add_paragraph()
        
        # Add course content
        self._add_course_content_to_docx(doc)
        
        # Add slide mapping
        self._add_slide_mapping_to_docx(doc)
        
        # Save document
        doc.save(output_path)
        print(f"📄 Course exported to: {output_path}")
        
        return output_path
    
    def _add_course_content_to_docx(self, doc):
        """Add course content sections to the DOCX document."""
        content_heading = doc.add_heading("📚 Course Content", level=1)
        
        for i, section in enumerate(self.content.sections, 1):
            self._add_content_section_to_docx(doc, section, level=2, section_number=i)
    
    def _add_content_section_to_docx(self, doc, section: ContentSection, level: int = 2, section_number: Optional[int] = None):
        """Recursively add ContentSection to DOCX document."""
        # Determine heading level (Word supports up to 9 heading levels)
        heading_level = min(level, 9)
        
        # Create section title
        title_prefix = f"{section_number}. " if section_number and level == 2 else ""
        section_heading = doc.add_heading(f"{title_prefix}{section.title}", level=heading_level)
        
        # Add content paragraphs
        if section.content:
            for paragraph_text in section.content:
                paragraph = doc.add_paragraph()
                # Add a bullet point for each paragraph
                paragraph.style = 'List Bullet'
                paragraph.add_run(paragraph_text)
        
        # Add subsections
        if section.subsections:
            for subsection in section.subsections:
                self._add_content_section_to_docx(doc, subsection, level + 1)
    
    def _add_slide_mapping_to_docx(self, doc):
        """Add slide mapping section to the DOCX document."""
        from collections import defaultdict
        
        # Add section separator
        doc.add_page_break()
        
        # Add slide mapping heading
        mapping_heading = doc.add_heading("📄 Slide Mapping", level=1)
        
        # Group slides by path for better readability
        path_to_slides = defaultdict(list)
        
        for mapping in self.content.slide_mappings:
            path_key = " → ".join(mapping.section_path) if mapping.section_path else "(unmapped)"
            path_to_slides[path_key].append(mapping.slide_number)
        
        # Create a table for better formatting
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set table headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Slides'
        hdr_cells[1].text = 'Section Path'
        
        # Add mapping data
        for path, slides in sorted(path_to_slides.items()):
            row_cells = table.add_row().cells
            slides_sorted = sorted(slides)
            slides_str = self._format_slide_ranges(slides_sorted)
            row_cells[0].text = slides_str
            row_cells[1].text = path
        
        # Add some spacing after the table
        doc.add_paragraph()
        
        # Add summary information
        summary_paragraph = doc.add_paragraph()
        summary_paragraph.add_run("Summary: ").bold = True
        summary_paragraph.add_run(f"Total slides mapped: {len(self.content.slide_mappings)}, "
                                 f"Unique sections: {len(path_to_slides)}")
    
    def _format_slide_ranges(self, slides: List[int]) -> str:
        """Format slide numbers into ranges (e.g., '1-3, 7, 9-11')."""
        if not slides:
            return ""
        
        ranges = []
        start = slides[0]
        end = slides[0]
        
        for slide in slides[1:]:
            if slide == end + 1:
                end = slide
            else:
                if start == end:
                    ranges.append(str(start))
                else:
                    ranges.append(f"{start}-{end}")
                start = slide
                end = slide
        
        # Add the last range
        if start == end:
            ranges.append(str(start))
        else:
            ranges.append(f"{start}-{end}")
        
        return ", ".join(ranges)


